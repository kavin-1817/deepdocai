import streamlit as st
import os
import tempfile
import re
import time
import random
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document  # For Word files
import openpyxl  # For Excel files
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
import google.generativeai as genai
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set up Streamlit page
st.set_page_config(page_title="DeepDocAI", page_icon="🤖", layout="wide")

# Load environment variables
try:
    load_dotenv()
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise ValueError("GOOGLE_API_KEY not found in environment variables")
    genai.configure(api_key=api_key)
except Exception as e:
    logger.error(f"Error configuring Google API: {str(e)}")
    st.error(f"Error configuring Google API: {str(e)}")

# Custom CSS with updated button effects
st.markdown("""
    <style>
    h1, .stHeader { border-bottom: none !important; }
    
    /* User Chat Bubble */
    .chat-bubble {
        background: linear-gradient(135deg, #9CA3AF 0%, #4B5563 100%);
        color: white;
        padding: 14px 20px;
        border-radius: 20px 20px 20px 5px;
        max-width: 80%;
        margin: 12px 0 !important;
        display: inline-block;
        font-size: 18px;
        line-height: 1.5;
        position: relative;
        box-shadow: 3px 3px 15px rgba(0, 0, 0, 0.2), inset 0 2px 5px rgba(255, 255, 255, 0.3);
        animation: fadeIn 0.5s ease-in-out, float 0.5s ease-in-out;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .chat-bubble::after {
        content: '';
        position: absolute;
        bottom: -8px;
        right: 10px;
        width: 0;
        height: 0;
        border-left: 10px solid transparent;
        border-right: 10px solid transparent;
        border-top: 10px solid #4B5563;
        transform: rotate(-20deg);
    }
    .chat-bubble:hover {
        transform: translateY(-3px);
        box-shadow: 5px 5px 20px rgba(0, 0, 0, 0.25);
    }
    
    /* AI Chat Bubble */
    .ai-bubble {
        background: linear-gradient(135deg, #A3BFFA 0%, #4B5EAA 100%);
        color: white;
        padding: 14px 20px;
        border-radius: 20px 20px 5px 20px;
        max-width: 80%;
        margin: 12px 0 !important;
        display: inline-block;
        font-size: 18px;
        line-height: 1.6;
        position: relative;
        box-shadow: 3px 3px 15px rgba(0, 0, 0, 0.2), inset 0 2px 5px rgba(255, 255, 255, 0.3);
        animation: fadeIn 0.5s ease-in-out, float 0.5s ease-in-out;
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .ai-bubble::after {
        content: '';
        position: absolute;
        bottom: -8px;
        left: 10px;
        width: 0;
        height: 0;
        border-left: 10px solid transparent;
        border-right: 10px solid transparent;
        border-top: 10px solid #4B5EAA;
        transform: rotate(20deg);
    }
    .ai-bubble:hover {
        transform: translateY(-3px);
        box-shadow: 5px 5px 20px rgba(0, 0, 0, 0.25);
    }
    
    /* Bold Text Styling */
    strong {
        font-weight: 700;
        text-shadow: 0.5px 0.5px 1px rgba(0, 0, 0, 0.2);
    }
    
    /* Sidebar Styling */
    .css-1v0mbdj, .css-1v3fvcr {
        background: linear-gradient(180deg, #9CA3AF 0%, #A3BFFA 100%);
        color: #1F2937;
        padding: 20px;
    }
    .css-1v0mbdj h1, .css-1v3fvcr h1, .css-1v0mbdj h2, .css-1v3fvcr h2 {
        color: #1F2937;
    }
    .stButton > button {
        background-color: #4B5EAA;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        font-size: 16px;
        transition: background-color 0.2s ease;
    }
    .stButton > button:hover {
        background-color: #4B5563;
    }
    .stTextInput > div > div > input {
    background-color: transparent;
    color: inherit; /* Uses the parent element's text color */
    border: none; /* Removes the border entirely */
    border-radius: 20px;
    padding: 8px 40px 8px 10px !important;
}
    .submit-button {
        color: #4B5EAA !important;
    }
    .submit-button:hover {
        color: #4B5563 !important;
    }
    .mic-button {
        background: none;
        border: none;
        cursor: pointer;
        font-size: 16px;
        color: #4B5563;
    }
    .mic-button.listening {
        color: #4B5EAA;
        animation: pulse 1s infinite;
    }
    
    /* Navigation Buttons Styling */
    .nav-buttons .stButton > button { /* Previous and Next buttons */
        background: linear-gradient(135deg, #9CA3AF 0%, #4B5563 100%);
        color: white;
        padding: 10px 20px;
        border: none;
        border-radius: 15px;
        font-size: 16px;
        min-width: 120px;
        box-shadow: 3px 3px 15px rgba(0, 0, 0, 0.2), inset 0 2px 5px rgba(255, 255, 255, 0.3);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .nav-buttons .stButton > button:hover:not(:disabled) {
        transform: translateY(-3px);
        box-shadow: 5px 5px 20px rgba(0, 0, 0, 0.25);
    }
    .nav-buttons .stButton > button:disabled {
        background: #D1D5DB;
        color: #6B7280;
        box-shadow: none;
        cursor: not-allowed;
    }
    .nav-buttons div { /* Page indicator */
        background: linear-gradient(135deg, #A3BFFA 0%, #4B5EAA 100%);
        color: white;
        padding: 10px 20px;
        border-radius: 15px;
        font-size: 16px;
        box-shadow: 3px 3px 15px rgba(0, 0, 0, 0.2), inset 0 2px 5px rgba(255, 255, 255, 0.3);
        transition: transform 0.2s ease, box-shadow 0.2s ease;
    }
    .nav-buttons div:hover {
        transform: translateY(-3px);
        box-shadow: 5px 5px 20px rgba(0, 0, 0, 0.25);
    }
    
    /* Animations */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(15px); }
        to { opacity: 1; transform: translateY(0); }
    }
    @keyframes float {
        0%, 100% { transform: translateY(0); }
        50% { transform: translateY(-5px); }
    }
    .typing { font-size: 14px; color: #888; animation: blink 1.5s infinite; }
    @keyframes blink { 0% { opacity: 0.2; } 50% { opacity: 1; } 100% { opacity: 0.2; } }
    .ai-response { animation: slideIn 0.5s ease-in-out; }
    @keyframes slideIn { from { opacity: 0; transform: translateX(-20px); } to { opacity: 1; transform: translateX(0); } }
    .user-response { animation: slideInRight 0.5s ease-in-out; }
    @keyframes slideInRight { from { opacity: 0; transform: translateX(20px); } to { opacity: 1; transform: translateX(0); } }
    .message-container { text-align: left; width: 100%; max-width: 800px; margin: 0px !important; padding: 0px !important; overflow: hidden; }
    .pagination-container { min-height: 400px; display: flex; flex-direction: column; justify-content: flex-start; align-items: center; padding: 0px !important; margin: 0px !important; overflow: hidden; }
    .nav-buttons { display: flex; justify-content: space-between; align-items: center; margin-top: 20px; width: 100%; max-width: 800px; }
    .stChatMessage { display: none !important; }
    .sidebar-input-container { background-color: #fff; border: 1px solid #4B5563; border-radius: 20px; padding: 10px; box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1); display: flex; align-items: center; gap: 8px; margin-top: 0px !important; margin-bottom: 0px !important; }
    .input-wrapper { position: relative; flex-grow: 1; }
    .stTextInput > div > div > input { border-radius: 20px; }
    @keyframes pulse { 0% { transform: scale(1); } 50% { transform: scale(1.2); } 100% { transform: scale(1); } }
    .stApp { margin: 0px !important; padding: 0px !important; }
    </style>
""", unsafe_allow_html=True)

# File Processing Functions
def extract_tables_from_pdf(pdf_path):
    tables_text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        tables_text += "| " + " | ".join(str(cell) for cell in table[0]) + " |\n"
                        tables_text += "| " + " | ".join(["---"] * len(table[0])) + " |\n"
                        for row in table[1:]:
                            tables_text += "| " + " | ".join(str(cell) for cell in row) + " |\n"
                        tables_text += "\n"
    except Exception as e:
        logger.error(f"Error extracting tables from PDF: {str(e)}")
        st.error(f"Error extracting tables from PDF: {str(e)}")
    return tables_text

def process_pdf(pdf):
    logger.info("Starting PDF processing")
    text = ""
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(pdf.getbuffer())
            temp_path = temp_file.name
        logger.info("PDF temp file created: %s", temp_path)

        pdf_reader = PdfReader(temp_path)
        for page in pdf_reader.pages:
            extracted_text = page.extract_text()
            if extracted_text:
                text += extracted_text + "\n\n"

        tables_text = extract_tables_from_pdf(temp_path)
        if tables_text.strip():
            text += "**Extracted Tables:**\n" + tables_text + "\n\n"
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        st.error(f"Error processing PDF: {str(e)}")
    finally:
        if temp_path and os.path.exists(temp_path):
            os.remove(temp_path)
    logger.info("PDF processing completed")
    return text

def process_docx(docx):
    logger.info("Starting DOCX processing")
    text = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            temp_file.write(docx.getbuffer())
            temp_path = temp_file.name
            doc = Document(temp_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n\n"
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    table_text += row_text + "\n"
                if table_text.strip():
                    text += "**Extracted Table from Word:**\n" + table_text + "\n\n"
            os.remove(temp_path)
    except Exception as e:
        logger.error(f"Error processing Word file: {str(e)}")
        st.error(f"Error processing Word file: {str(e)}")
    return text

def process_xlsx(xlsx):
    logger.info("Starting XLSX processing")
    text = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp_file:
            temp_file.write(xlsx.getbuffer())
            temp_path = temp_file.name
            wb = openpyxl.load_workbook(temp_path)
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"**Sheet: {sheet_name}**\n"
                for row in sheet.rows:
                    row_text = " | ".join(str(cell.value) if cell.value is not None else "" for cell in row)
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            os.remove(temp_path)
    except Exception as e:
        logger.error(f"Error processing Excel file: {str(e)}")
        st.error(f"Error processing Excel file: {str(e)}")
    return text

def get_file_text(docs):
    logger.info("Extracting text from files")
    text = ""
    if not docs:
        return text
    try:
        from concurrent.futures import ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=4) as executor:
            futures = []
            for doc in docs:
                if doc.name.endswith('.pdf'):
                    futures.append(executor.submit(process_pdf, doc))
                elif doc.name.endswith('.docx'):
                    futures.append(executor.submit(process_docx, doc))
                elif doc.name.endswith('.xlsx'):
                    futures.append(executor.submit(process_xlsx, doc))
            for future in futures:
                text += future.result() + "\n"
    except Exception as e:
        logger.error(f"Processing failed: {str(e)}")
        st.error(f"Processing failed: {str(e)}")
    logger.info("Text extraction completed: %d characters", len(text))
    return text

def get_text_chunks(text):
    try:
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=5000, chunk_overlap=500)
        return text_splitter.split_text(text)
    except Exception as e:
        logger.error(f"Error splitting text: {str(e)}")
        st.error(f"Error splitting text: {str(e)}")
        return []

def get_vector_store(text_chunks):
    if not text_chunks:
        logger.error("No text found in the uploaded files")
        st.error("❌ No text found in the uploaded files.")
        return False
    try:
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
        vector_store.save_local("faiss_index")
        logger.info("Vector store created successfully")
        return True
    except Exception as e:
        logger.error(f"Error in vector storage: {str(e)}")
        st.error(f"⚠️ Error in vector storage: {e}")
        return False

# AI Interaction Functions
prompt_template = """
You are an advanced AI assistant with strong reasoning capabilities. Your task is to provide **clear, well-structured, and fact-based answers** strictly using the given context.

### **Guidelines:**  
1. If the answer **exists in the context**, provide a **concise, direct, and well-explained response**.  
2. If the answer **is not found directly**, analyze the context carefully:  
   - If the context provides indirect hints, use reasoning to give the most accurate answer possible.  
   - If the context explicitly states the **opposite** of what is being asked, state that clearly.  
   - If the answer is **completely unavailable**, respond in a human-like way by saying:  
     **"No, [subject] does not [action]."**  
     or  
     **"There is no information available in the provided context to determine this."**  
3. **Use bullet points and line breaks for clarity when listing multiple points.**  
4. **Never make up information that is not in the context.**  

### **Context:**  
{context}  

### **User Question:**  
{question}  

### **Answer:**  
"""

def get_conversational_chain():
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    model = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

def format_response(response_text):
    formatted_text = re.sub(r'-\n', '', response_text)
    formatted_text = re.sub(r'^\s*\*+\s+', '• ', formatted_text, flags=re.MULTILINE)
    formatted_text = formatted_text.replace("\n-", "\n\n• ")
    formatted_text = formatted_text.replace("-", " - ")
    formatted_text = re.sub(r'\n\s*\n+', '\n\n', formatted_text)
    formatted_text = "\n\n".join([line.strip() for line in formatted_text.split("\n") if line.strip()])
    return formatted_text

def display_response(response_text, chat_placeholder):
    formatted_text = format_response(response_text)
    with chat_placeholder:
        st.markdown(
            f'<div class="message-container"><div class="ai-bubble ai-response">{formatted_text}</div></div>',
            unsafe_allow_html=True
        )

def user_input(user_question, chat_placeholder):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local("faiss_index", embeddings, allow_dangerous_deserialization=True)

    docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()

    with chat_placeholder:
        st.markdown(
            f'<div class="message-container"><div class="chat-bubble user-response">🧑‍💼 You: {user_question}</div></div>',
            unsafe_allow_html=True
        )

    response = chain(
        {"input_documents": docs, "question": user_question},
        return_only_outputs=True
    )

    response_text = format_response(response["output_text"])
    
    display_response(response_text, chat_placeholder)
    st.session_state.conversation.append({"role": "ai", "content": response_text})

def main():
    logger.info("App started")
    if "conversation" not in st.session_state:
        st.session_state.conversation = []
    if "current_page" not in st.session_state:
        st.session_state.current_page = 0
    if "processed" not in st.session_state:
        st.session_state.processed = False
    if "docs" not in st.session_state:
        st.session_state.docs = []

    chat_placeholder = st.container()

    if st.session_state.conversation:
        message_pairs = []
        for i in range(0, len(st.session_state.conversation), 2):
            pair = [st.session_state.conversation[i]]
            if i + 1 < len(st.session_state.conversation):
                pair.append(st.session_state.conversation[i + 1])
            message_pairs.append(pair)
        
        total_pages = len(message_pairs)
        current_page = st.session_state.current_page
        
        if total_pages > 0:
            with chat_placeholder:
                st.markdown('<div class="message-container">', unsafe_allow_html=True)
                for message in message_pairs[current_page]:
                    if message["role"] == "user":
                        st.markdown(
                            f'<div class="chat-bubble user-response">🧑‍💼 You: {message["content"]}</div>',
                            unsafe_allow_html=True
                        )
                    else:
                        st.markdown(
                            f'<div class="ai-bubble ai-response">{message["content"]}</div>',
                            unsafe_allow_html=True
                        )
                st.markdown('</div>', unsafe_allow_html=True)
            
            with chat_placeholder:
                st.markdown('<div class="nav-buttons">', unsafe_allow_html=True)
                col1, col2, col3 = st.columns([1, 1, 1])
                
                with col1:
                    if st.button("← Previous", key="prev", disabled=current_page == 0, use_container_width=True):
                        st.session_state.current_page = max(0, current_page - 1)
                        st.rerun()
                
                with col2:
                    st.markdown(f"<div style='text-align: center;'>Page {current_page + 1} of {total_pages}</div>", unsafe_allow_html=True)
                
                with col3:
                    if st.button("Next →", key="next", disabled=current_page == total_pages - 1, use_container_width=True):
                        st.session_state.current_page = min(total_pages - 1, current_page + 1)
                        st.rerun()
                
                st.markdown('</div>', unsafe_allow_html=True)

    with st.sidebar:
        st.title("🚀 DeepDocAI")
        st.subheader("📂 Upload Documents:")
        docs = st.file_uploader(
            "Upload PDFs, Word (.docx), or Excel (.xlsx) files and Click on Submit",
            accept_multiple_files=True,
            type=['pdf', 'docx', 'xlsx'],
            key="file_uploader"
        )
        
        if st.button("📥 Submit & Process", key="submit_button"):
            logger.info("Submit button clicked")
            if not docs:
                st.warning("⚠️ Please upload at least one file")
            else:
                st.session_state.docs = docs
                logger.info("Processing files")
                raw_text = get_file_text(docs)
                if not raw_text.strip():
                    st.error("❌ Failed to extract text from files.")
                else:
                    text_chunks = get_text_chunks(raw_text)
                    if get_vector_store(text_chunks):
                        st.session_state.processed = True

        with st.form(key="input_form", clear_on_submit=True):
            col1, col2 = st.columns([12, 1])
            with col1:
                user_question = st.text_input(
                    "Ask a question...",
                    placeholder="Type your question here...",
                    label_visibility="collapsed",
                    value=""
                )
            with col2:
                submit_button = st.form_submit_button("➤", use_container_width=False)

    if submit_button and user_question:
        st.session_state.conversation.append({"role": "user", "content": user_question})
        user_input(user_question, chat_placeholder)
        st.session_state.current_page = (len(st.session_state.conversation) // 2) - 1
        st.rerun()

    logger.info("Main loop completed")

if __name__ == "__main__":
    main()
